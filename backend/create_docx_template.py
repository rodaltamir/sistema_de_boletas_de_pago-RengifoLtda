import docx

def create_template():
    doc = docx.Document()
    
    doc.add_heading('PRELIQUIDACIÓN O PREFINIQUITO', 0)
    
    doc.add_paragraph('NOMBRE DEL TRABAJADOR: {{ nombre_trabajador }}')
    doc.add_paragraph('RAZÓN SOCIAL DEL EMPLEADOR: {{ razon_social }}')
    doc.add_paragraph('FECHA DE INGRESO: {{ fecha_ingreso }}')
    doc.add_paragraph('FECHA DE RETIRO: {{ fecha_retiro }}')
    doc.add_paragraph('TIEMPO DE TRABAJO: {{ anios }} Años, {{ meses }} Meses, {{ dias }} Días')
    doc.add_paragraph('SUELDO PROMEDIO (Bs): {{ sueldo_promedio }}')
    
    doc.add_heading('BENEFICIOS SOCIALES', level=1)
    doc.add_paragraph('DESAHUCIO (3 meses): {{ desahucio }}')
    doc.add_paragraph('INDEMNIZACIÓN: {{ indemnizacion }}')
    doc.add_paragraph('AGUINALDO: {{ aguinaldo }}')
    doc.add_paragraph('VACACIONES: {{ vacaciones }}')
    
    doc.add_heading('DESCUENTOS', level=1)
    doc.add_paragraph('TOTAL DESCUENTOS: {{ total_descuentos }}')
    
    doc.add_heading('TOTALES', level=1)
    doc.add_paragraph('TOTAL CÁLCULO: {{ total_calculo }}')
    doc.add_paragraph('MULTA 30%: {{ multa }}')
    doc.add_paragraph('TOTAL FINAL: {{ total_final }}')
    
    doc.save('app/templates/prefiniquito_template.docx')
    print("Plantilla Word creada exitosamente.")

if __name__ == "__main__":
    create_template()
